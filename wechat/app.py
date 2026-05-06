import os
import io
from flask import Flask, request, render_template_string, send_file
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor
from bs4 import BeautifulSoup

app = Flask(__name__)

client = OpenAI(
    api_key=os.environ.get("DEEPSEEK_API_KEY"),
    base_url="https://api.deepseek.com"
)

HTML_PAGE = """
<!DOCTYPE html>
<html lang="zh">
<head>
<meta charset="UTF-8">
<title>公众号排版助手</title>
<style>
  body { font-family: Arial, sans-serif; max-width: 900px; margin: 40px auto; padding: 20px; background: #f5f5f5; }
  h2 { color: #07c160; }
  textarea { width: 100%; height: 200px; padding: 10px; font-size: 14px; border-radius: 5px; border: 1px solid #ddd; }
  .btn-group { margin-top: 10px; display: flex; gap: 10px; flex-wrap: wrap; }
  button { color: white; border: none; padding: 12px 30px; font-size: 16px; cursor: pointer; border-radius: 5px; }
  #convertBtn { background: #07c160; }
  #copyBtn { background: #1890ff; display: none; }
  #wordBtn { background: #fa8c16; display: none; }
  #result { margin-top: 30px; background: white; border: 1px solid #ddd; padding: 20px; min-height: 100px; border-radius: 5px; }
</style>
</head>
<body>
<h2>📝 公众号自动排版助手</h2>
<textarea id="input" placeholder="在这里粘贴你的文章内容..."></textarea>
<div class="btn-group">
  <button id="convertBtn" onclick="convert()">🚀 一键排版</button>
  <button id="copyBtn" onclick="copyHTML()">📋 复制HTML</button>
  <button id="wordBtn" onclick="exportWord()">📄 导出Word</button>
</div>
<div id="result"><p style="color:#999">排版结果将在这里显示...</p></div>
<script>
let generatedHTML = "";
async function convert() {
  const text = document.getElementById("input").value;
  if (!text.trim()) { alert("请先输入文章内容"); return; }
  document.getElementById("result").innerHTML = "<p>⏳ AI排版中，请稍候...</p>";
  const res = await fetch("/convert", {
    method: "POST",
    headers: {"Content-Type": "application/json"},
    body: JSON.stringify({text})
  });
  const data = await res.json();
  generatedHTML = data.html;
  document.getElementById("result").innerHTML = generatedHTML;
  document.getElementById("copyBtn").style.display = "inline-block";
  document.getElementById("wordBtn").style.display = "inline-block";
}
function copyHTML() {
  navigator.clipboard.writeText(generatedHTML);
  alert("✅ HTML已复制！去公众号编辑器粘贴吧");
}
async function exportWord() {
  document.getElementById("wordBtn").innerText = "⏳ 生成中...";
  const res = await fetch("/export_word", {
    method: "POST",
    headers: {"Content-Type": "application/json"},
    body: JSON.stringify({html: generatedHTML})
  });
  const blob = await res.blob();
  const url = window.URL.createObjectURL(blob);
  const a = document.createElement("a");
  a.href = url;
  a.download = "文章.docx";
  a.click();
  document.getElementById("wordBtn").innerText = "📄 导出Word";
}
</script>
</body>
</html>
"""

@app.route("/")
def index():
    return render_template_string(HTML_PAGE)

@app.route("/convert", methods=["POST"])
def convert():
    text = request.json.get("text", "")
    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "system", "content": "你是微信公众号编辑。将文章转换为排版精美的HTML。要求：标题用<h2>加粗居中，段落用<p>行距1.8，小标题用<h3>，重点内容用<strong>。只输出HTML代码片段，不要<!DOCTYPE>和<html>标签。"},
            {"role": "user", "content": text}
        ]
    )
    html = response.choices[0].message.content
    return {"html": html}

@app.route("/export_word", methods=["POST"])
def export_word():
    html = request.json.get("html", "")
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()

    for element in soup.find_all(["h1","h2","h3","p","strong"]):
        tag = element.name
        text = element.get_text()
        if not text.strip():
            continue
        if tag in ["h1", "h2"]:
            p = doc.add_heading(text, level=1)
            p.runs[0].font.color.rgb = RGBColor(0x07, 0xC1, 0x60)
        elif tag == "h3":
            p = doc.add_heading(text, level=2)
        else:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(12)
            if tag == "strong":
                run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name="文章.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)